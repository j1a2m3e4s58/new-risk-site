from django.conf import settings
from django.contrib import messages
from django.contrib.auth.decorators import login_required, permission_required
from django.db.models import Q
from django.http import FileResponse, HttpResponseForbidden
from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse
from django.utils import timezone
from django.http import HttpResponse
from datetime import timedelta
import csv
import os
import re
import shutil
import tempfile
import zipfile
from io import BytesIO
from urllib.parse import urlencode
from .models import CustomerRiskProfile, RiskAssessment, ReportConfiguration, RiskIncident, SystemAuditLog
from .forms import RiskActionUpdateForm, RiskAssessmentForm, RiskIncidentForm

# ========= ZERO_OCCURRENCE_HELPER_START =========
def is_zero_occurrence(value) -> bool:
    if value is None:
        return True

    v = str(value).strip().lower()

    ZERO_WORDS = [
        "0", "0.0", "zero", "nil", "none", "no", "n/a", "",
        "always updated", "timelines met", "on time", "no issues", "ok"
    ]

    return v in ZERO_WORDS
# ========= ZERO_OCCURRENCE_HELPER_END =========


# ========= UNIQUE_ID_GLOBAL_START =========
def make_unique_reference_id(base_ref):
    ref = base_ref
    bump = 1
    while RiskAssessment.objects.filter(reference_id=ref).exists():
        ref = f"{base_ref}-{bump}"
        bump += 1
    return ref
# ========= UNIQUE_ID_GLOBAL_END =========


def log_system_event(request, action, summary, target=None, metadata=None, area_name=""):
    user = getattr(request, "user", None)
    if not getattr(user, "is_authenticated", False) or not getattr(user, "pk", None):
        user = None

    SystemAuditLog.objects.create(
        user=user,
        action=action,
        target_model=target.__class__.__name__ if target else "",
        target_id=str(getattr(target, "pk", "")) if target else "",
        reference_id=getattr(target, "reference_id", "") if target else "",
        area_name=area_name or getattr(target, "area_name", "") or "",
        summary=summary,
        metadata=metadata or {},
    )


RISK_LEVELS = ["High", "Medium", "Low"]
HEATMAP_PROBABILITIES = ["High", "Medium", "Low"]
HEATMAP_IMPACTS = ["Low", "Medium", "High"]


def normalize_risk_level(value):
    value = (value or "").strip()
    if value in ["Very High", "High", "Critical"]:
        return "High"
    if value in ["Medium", "Moderate", "Severe"]:
        return "Medium"
    return "Low"


def display_level(value):
    return normalize_risk_level(value)


def risk_rank_value(value):
    return {"High": 3, "Medium": 2, "Low": 1}.get(normalize_risk_level(value), 1)


def is_high_risk(value):
    return normalize_risk_level(value) == "High"


def is_medium_or_high_risk(value):
    return normalize_risk_level(value) in ["Medium", "High"]


def _in_group(user, *group_names):
    return bool(
        getattr(user, "is_authenticated", False)
        and user.groups.filter(name__in=group_names).exists()
    )


def can_submit_risk(user):
    return getattr(user, "is_authenticated", False) and not _in_group(user, "Board")


def can_review_risk(user):
    return getattr(user, "is_staff", False) or _in_group(user, "Risk Officer", "Admin")


def can_approve_delete(user):
    return getattr(user, "is_superuser", False) or _in_group(user, "Admin")


def can_view_reports(user):
    return getattr(user, "is_authenticated", False)


def _department_score(area_name, risks):
    risk_list = list(risks)
    high = sum(1 for risk in risk_list if is_high_risk(risk.residual_rating))
    medium = sum(1 for risk in risk_list if normalize_risk_level(risk.residual_rating) == "Medium")
    overdue = sum(1 for risk in risk_list if risk.is_action_overdue)
    escalated = sum(1 for risk in risk_list if risk.escalation_status != "Normal")
    open_actions_count = sum(
        1 for risk in risk_list
        if risk.mitigation_action and risk.action_status != "Completed"
    )
    return {
        "area": area_name,
        "score": (high * 20) + (medium * 10) + (overdue * 10) + (escalated * 10) + (open_actions_count * 3) + len(risk_list),
        "total": len(risk_list),
        "high": high,
        "medium": medium,
        "overdue": overdue,
        "escalated": escalated,
        "open_actions": open_actions_count,
    }


def _rating_chart(risks):
    risk_list = list(risks)
    rows = []
    for rating in RISK_LEVELS:
        count = sum(1 for risk in risk_list if risk.residual_rating == rating)
        rows.append({
            "label": rating,
            "count": count,
            "percent": round((count / len(risk_list)) * 100) if risk_list else 0,
        })
    return rows


def _apply_dashboard_filters(request, risks):
    q = request.GET.get("q", "").strip()
    owner = request.GET.get("owner", "").strip()
    rating = request.GET.get("rating", "").strip()
    action_status = request.GET.get("action_status", "").strip()
    escalation = request.GET.get("escalation", "").strip()
    due = request.GET.get("due", "").strip()

    if q:
        risks = risks.filter(
            Q(reference_id__icontains=q)
            | Q(description__icontains=q)
            | Q(area_name__icontains=q)
            | Q(risk_owner__icontains=q)
            | Q(control_owner__icontains=q)
            | Q(action_responsible_officer__icontains=q)
        )
    if owner:
        risks = risks.filter(risk_owner__icontains=owner)
    if rating:
        risks = risks.filter(residual_rating=rating)
    if action_status:
        risks = risks.filter(action_status=action_status)
    if escalation:
        risks = risks.filter(escalation_status=escalation)
    if due == "overdue":
        today = timezone.localdate()
        risks = risks.filter(action_due_date__lt=today).exclude(action_status="Completed")
    elif due == "next30":
        today = timezone.localdate()
        risks = risks.filter(action_due_date__gte=today, action_due_date__lte=today + timedelta(days=30))
    elif due:
        risks = risks.filter(action_due_date=due)
    return risks


def _build_assurance_cockpit(risks, incidents, departments):
    risk_list = list(risks)
    incident_list = list(incidents)
    today = timezone.localdate()
    total = len(risk_list) or 1

    missing_owner = sum(1 for risk in risk_list if not (risk.risk_owner or "").strip())
    missing_controls = sum(1 for risk in risk_list if not (risk.controls or "").strip())
    missing_actions = sum(
        1 for risk in risk_list
        if is_medium_or_high_risk(risk.residual_rating) and not (risk.mitigation_action or "").strip()
    )
    stale_reviews = sum(1 for risk in risk_list if risk.updated_at.date() < today - timedelta(days=90))
    overdue = sum(1 for risk in risk_list if risk.is_action_overdue)
    weak_controls = sum(1 for risk in risk_list if risk.control_effectiveness == "Weak")
    high = sum(1 for risk in risk_list if is_high_risk(risk.residual_rating))
    board_items = sum(1 for risk in risk_list if risk.escalation_status == "Board Attention")
    open_incidents = sum(1 for incident in incident_list if incident.status not in ["Resolved", "Closed"])
    pending_approvals = sum(1 for risk in risk_list if risk.workflow_status in ["Draft", "Reviewed"])
    action_progress = round(sum(risk.action_progress or 0 for risk in risk_list) / total)
    due_soon = sum(
        1 for risk in risk_list
        if risk.action_due_date
        and today <= risk.action_due_date <= today + timedelta(days=30)
        and risk.action_status != "Completed"
    )

    quality_penalty = missing_owner + missing_controls + missing_actions + stale_reviews
    quality_score = max(0, round(100 - ((quality_penalty / max(total * 4, 1)) * 100)))
    appetite_status = "Breach" if high or board_items else "Within appetite"
    assurance_items = [
        {"label": "Data quality", "value": f"{quality_score}%", "tone": "good" if quality_score >= 75 else "warn"},
        {"label": "Appetite status", "value": appetite_status, "tone": "bad" if appetite_status == "Breach" else "good"},
        {"label": "High exposure", "value": high, "tone": "bad" if high else "good"},
        {"label": "Board attention", "value": board_items, "tone": "bad" if board_items else "good"},
        {"label": "Overdue actions", "value": overdue, "tone": "bad" if overdue else "good"},
        {"label": "Due in 30 days", "value": due_soon, "tone": "warn" if due_soon else "good"},
        {"label": "Avg action progress", "value": f"{action_progress}%", "tone": "good" if action_progress >= 70 else "warn"},
        {"label": "Weak controls", "value": weak_controls, "tone": "bad" if weak_controls else "good"},
        {"label": "Missing controls", "value": missing_controls, "tone": "warn" if missing_controls else "good"},
        {"label": "Missing owners", "value": missing_owner, "tone": "warn" if missing_owner else "good"},
        {"label": "No action plan", "value": missing_actions, "tone": "bad" if missing_actions else "good"},
        {"label": "Stale reviews", "value": stale_reviews, "tone": "warn" if stale_reviews else "good"},
        {"label": "Open incidents", "value": open_incidents, "tone": "bad" if open_incidents else "good"},
        {"label": "Departments covered", "value": len(departments), "tone": "good" if departments else "warn"},
        {"label": "Pending approvals", "value": pending_approvals, "tone": "warn" if pending_approvals else "good"},
    ]
    return assurance_items


# ========= RISK_OWNER_SUGGEST_START =========
def suggest_risk_owner(area_name):
    a = (area_name or "").strip().lower()

    if "microfinance" in a:
        return "Head of Microfinance"
    if "credit" in a:
        return "Head of Credit"
    if "finance" in a:
        return "Head of Finance"
    if a == "it" or " ict" in f" {a} " or " it " in f" {a} " or "information technology" in a:
        return "Head of IT"
    if "operations" in a or "teller" in a or "customer service" in a:
        return "Head of Operations"
    if "compliance" in a:
        return "Compliance Officer"
    if "audit" in a:
        return "Internal Auditor"
    if "treasury" in a:
        return "Treasury Manager"
    if "hr" in a or "human resource" in a:
        return "Head of HR"
    if "legal" in a:
        return "Legal Officer"

    return "Department Head"
# ========= RISK_OWNER_SUGGEST_END =========


# ========= SMART_SCORING_START =========
def score_probability_from_occurrence(occurrence_value):
    """
    occurrence_value can be '', '0', ' 200', '10', etc.
    Returns one of: Low, Medium, High
    """
    try:
        n = int(str(occurrence_value).strip())
    except Exception:
        n = 0

    if n <= 0:
        return "Low"
    if n <= 2:
        return "Low"
    if n <= 5:
        return "Medium"
    return "High"


def score_impact_from_text(related_risk_text):
    """
    Keyword-based impact scoring from Related Risk / Description text.
    Returns: Low/Medium/High
    """
    t = (related_risk_text or "").lower()

    very_high_keys = [
        "robbery", "fraud", "theft", "pilfer", "unauthorized", "suppression",
        "money laundering", "aml", "cft", "penalty", "regulatory", "impersonation",
        "asset loss", "loss of funds", "e-money", "identity theft"
    ]
    high_keys = [
        "reputational", "customer complaint", "complaints", "data privacy", "information leakage",
        "service", "downtime"
    ]

    for k in very_high_keys:
        if k in t:
            return "High"

    for k in high_keys:
        if k in t:
            return "High"

    return "Medium"


def default_controls_for_area(area_name):
    a = (area_name or "").lower()
    if "teller" in a or "customer service" in a or "operations" in a:
        return "Maker-checker, daily call-over, cash limits, CCTV monitoring, ID verification"
    if "credit" in a or "microfinance" in a:
        return "Approval workflow controls, KYC verification, monitoring visits, collections follow-up"
    if "it" in a or "ict" in a:
        return "Access control, system monitoring, change management, alerting, backups"
    if "compliance" in a or "aml" in a:
        return "Transaction monitoring, reporting controls, periodic compliance review"
    return "Standard Controls"
# ========= SMART_SCORING_END =========


CUSTOMER_PROFILE_FACTORS = [
    ("banking_relationship", "Length of banking relationship", 1),
    ("kyc_exception", "KYC information exception", 1),
    ("high_value", "High net worth or high value transaction", 3),
    ("pep", "Politically exposed person / close associate", 1),
    ("complex_ownership", "Complex control or ownership structure", 3),
    ("source_of_funds", "Unclear or undocumented source of funds", 1),
    ("product_services", "Product and services risk", 3),
    ("delivery_channel", "Transaction / delivery channel risk", 1),
    ("location", "Nationality and country of residence", 1),
    ("watchlist", "Watch list filtering", 1),
]


def _split_extracted_row(line):
    if "\t" in line:
        return [p.strip() for p in line.split("\t") if p.strip()]
    return [p.strip() for p in re.split(r"\s{2,}", line.strip()) if p.strip()]


def _clean_ref_part(value):
    cleaned = re.sub(r"[^A-Z0-9\-]", "", (value or "").upper().replace(" ", "-"))
    return cleaned[:16] or "GENERAL"


def _profile_rating_from_average(average_score):
    if average_score <= 1.2:
        return "Low"
    if average_score <= 2:
        return "Medium"
    return "High"


def _profile_levels_from_rating(profile_rating):
    if profile_rating == "High":
        return "High", "High"
    if profile_rating == "Medium":
        return "Medium", "High"
    return "Low", "Low"


def _edd_recommendation(profile_rating, high_factors):
    if profile_rating == "High":
        return (
            "Enhanced due diligence required: verify source of funds, confirm beneficial ownership, "
            "screen related parties, apply closer transaction monitoring, obtain senior management approval, "
            "and schedule a shorter review cycle."
        )
    if high_factors:
        return (
            "Medium monitoring required: review the high-weight determinants, confirm KYC completeness, "
            "validate expected transaction behavior, and document management review."
        )
    return "Standard customer due diligence and periodic monitoring are appropriate based on the current profile."


def _profile_confidence_notes(source_type, determinant_rows, average_score, profile_rating):
    source_label = source_type or "pasted text"
    return (
        f"Source: {source_label}. Detected {len(determinant_rows)} determinant(s). "
        f"Applied template scale Low <= 1.20, Medium 1.21-2.00, High 2.01-3.00. "
        f"Average score {average_score} produced {profile_rating} profile rating."
    )


def _review_plan_for_rating(profile_rating):
    rating = (profile_rating or "").lower()
    if rating == "high":
        return "Quarterly / enhanced due diligence", timezone.localdate() + timedelta(days=90)
    if rating == "medium":
        return "Six-month review", timezone.localdate() + timedelta(days=180)
    return "Annual review", timezone.localdate() + timedelta(days=365)


def _profile_rank(profile_rating):
    return {"Low": 1, "Medium": 2, "High": 3}.get(normalize_risk_level(profile_rating), 0)


def _previous_customer_profile(account_no="", account_name=""):
    profiles = CustomerRiskProfile.objects.all()
    if account_no:
        match = profiles.filter(account_no=account_no).first()
        if match:
            return match
    if account_name:
        return profiles.filter(account_name__iexact=account_name).first()
    return None


def _profile_duplicate_context(profile_data):
    if not profile_data:
        return None
    previous = _previous_customer_profile(
        account_no=profile_data.get("account_no", ""),
        account_name=profile_data.get("account_name", ""),
    )
    if not previous:
        return None
    current_rating = profile_data.get("profile_rating", "")
    movement = "Unchanged"
    if _profile_rank(current_rating) > _profile_rank(previous.profile_rating):
        movement = "Worsened"
    elif _profile_rank(current_rating) < _profile_rank(previous.profile_rating):
        movement = "Improved"
    return {
        "id": previous.id,
        "previous_rating": previous.profile_rating,
        "previous_score": previous.average_score,
        "current_rating": current_rating,
        "current_score": profile_data.get("average_score"),
        "movement": movement,
        "created_at": previous.created_at,
    }


def _uploaded_file_to_text(uploaded_file):
    if not uploaded_file:
        return "", "", ""

    filename = uploaded_file.name or "uploaded-file"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    data = uploaded_file.read()

    if ext in ["xlsx", "xlsm"]:
        from openpyxl import load_workbook

        workbook = load_workbook(BytesIO(data), data_only=True)
        lines = []
        for worksheet in workbook.worksheets:
            lines.append(f"Sheet: {worksheet.title}")
            for row in worksheet.iter_rows(values_only=True):
                values = [str(value).strip() for value in row if value not in [None, ""]]
                if values:
                    lines.append("\t".join(values))
        return "\n".join(lines), filename, "excel"

    if ext == "docx":
        from docx import Document

        document = Document(BytesIO(data))
        lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip().replace("\n", " ") for cell in row.cells if cell.text.strip()]
                if values:
                    lines.append("\t".join(values))
        return "\n".join(lines), filename, "word"

    if ext == "pdf":
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(data))
        lines = []
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                lines.append(text.strip())
        return "\n".join(lines), filename, "pdf"

    text = data.decode("utf-8", errors="replace")
    return text, filename, "text"


def _extract_numeric_weight(parts):
    for part in reversed(parts):
        text = str(part).strip()
        if re.fullmatch(r"[123](?:\.0)?", text):
            return int(float(text))
    return None


def _weight_from_profile_text(label, value):
    text = f"{label} {value}".lower()

    if any(k in text for k in ["politically exposed", "pep", "sanction", "watch list match", "watchlist match"]):
        return 3
    if any(k in text for k in ["high net worth", "high value", "complex control", "complex ownership", "unclear source", "undocumented source", "source of funds"]):
        return 3
    if any(k in text for k in ["0-12", "0 - 12", "less than 12", "new customer"]):
        return 3
    if any(k in text for k in ["2-<4", "2 -", "2 years", "3 years", "medium"]):
        return 2
    if any(k in text for k in [">5", "more than 5", "ghana", "no match", "low", "current account", "mobile banking", "ussd"]):
        return 1
    return 1


def _parse_customer_risk_profile(raw_text, source_type="", source_filename=""):
    lines = [ln.strip() for ln in raw_text.splitlines() if ln.strip()]
    lowered = "\n".join(lines).lower()
    if "customer risk profile" not in lowered and "risk determinants" not in lowered:
        return None

    account_no = ""
    account_name = ""
    for i, line in enumerate(lines):
        parts = _split_extracted_row(line)
        labels = [p.lower().rstrip(":") for p in parts]
        if "account no" in " ".join(labels) and i + 1 < len(lines):
            next_parts = _split_extracted_row(lines[i + 1])
            if next_parts:
                account_no = next_parts[0]
            if len(next_parts) > 1:
                account_name = next_parts[1]
        if "name of account" in line.lower() and not account_name and len(parts) > 1:
            account_name = parts[1]

    determinant_rows = []
    current_category = ""
    in_determinants = False
    for line in lines:
        parts = _split_extracted_row(line)
        if len(parts) < 2:
            continue
        joined = " ".join(parts).lower()
        if "risk determinants" in joined or "risk variables" in joined:
            in_determinants = True
            continue
        if "total risk score" in joined:
            break
        if not in_determinants:
            continue
        if any(skip in joined for skip in ["average risk score", "rating", "scale", "prepared by", "reviewed by", "approved by", "sheet:"]):
            continue

        weight = _extract_numeric_weight(parts)
        if weight is None:
            continue

        category = parts[0] if len(parts) >= 3 else current_category
        variable = parts[1] if len(parts) >= 3 else parts[0]
        if category:
            current_category = category
        category = category or current_category or "Customer"
        determinant_rows.append({
            "category": category,
            "variable": variable,
            "weight": weight,
        })

    if not determinant_rows:
        inferred = []
        for key, label, default_weight in CUSTOMER_PROFILE_FACTORS:
            inferred_weight = _weight_from_profile_text(label, raw_text) or default_weight
            inferred.append({"category": "Customer Risk Profile", "variable": label, "weight": inferred_weight})
        determinant_rows = inferred

    total_score = sum(row["weight"] for row in determinant_rows)
    average_score = round(total_score / max(len(determinant_rows), 1), 2)
    profile_rating = _profile_rating_from_average(average_score)
    probability, impact = _profile_levels_from_rating(profile_rating)
    high_factors = [row for row in determinant_rows if row["weight"] >= 3]
    account_label = account_name or account_no or "Customer"
    high_factor_text = "; ".join(f"{row['category']}: {row['variable']}" for row in high_factors[:6]) or "No high-weight determinant identified."
    recommendation = _edd_recommendation(profile_rating, high_factors)
    confidence_notes = _profile_confidence_notes(source_type, determinant_rows, average_score, profile_rating)

    base_ref = f"CRP-{_clean_ref_part(account_label)}"
    reference_id = make_unique_reference_id(base_ref)
    rating_label = profile_rating

    return {
        "reference_id": reference_id,
        "area_name": "Customer Risk Profile",
        "reporting_period": "",
        "risk_owner": "Compliance Officer",
        "risk_coordinator_name": "Risk & Compliance Coordinator",
        "risk_description": f"Customer risk profile for {account_label}: {rating_label} customer risk.",
        "root_cause": f"Customer profile determinant average score is {average_score} based on {len(determinant_rows)} factor(s).",
        "trigger": "Customer risk profile assessment indicates elevated customer due diligence exposure.",
        "consequences": "Potential AML/CFT, KYC, transaction monitoring, reputational, and regulatory exposure if customer risk is not monitored.",
        "inherent_probability": probability,
        "inherent_impact": impact,
        "inherent_rating": "-",
        "control_descriptions": "Customer due diligence, beneficial ownership verification, watchlist screening, transaction monitoring, enhanced due diligence where required, periodic review.",
        "control_owner": "Compliance Officer",
        "residual_probability": probability,
        "residual_impact": normalize_risk_level(impact),
        "residual_rating": "-",
        "source_kri": "Customer Risk Profile",
        "source_kri_description": high_factor_text,
        "source_related_risk": f"Average score {average_score}; profile rating {rating_label}.",
        "source_process": "Customer due diligence / AML monitoring",
        "source_occurrence": str(total_score),
        "customer_profile_score": average_score,
        "customer_profile_rating": rating_label,
        "customer_profile_notes": "\n".join(f"{row['category']} | {row['variable']} | Weight {row['weight']}" for row in determinant_rows),
        "customer_profile_data": {
            "account_no": account_no,
            "account_name": account_name,
            "profile_rating": rating_label,
            "average_score": average_score,
            "total_score": total_score,
            "determinant_count": len(determinant_rows),
            "determinants": determinant_rows,
            "recommendation": recommendation,
            "enhanced_due_diligence": recommendation,
            "confidence_notes": confidence_notes,
            "source_filename": source_filename,
            "source_type": source_type,
        },
        "recommendation": recommendation,
        "confidence_notes": confidence_notes,
    }


def _likelihood_from_occurrence(value):
    v = str(value or "").strip().lower()

    if not v:
        return "Medium"
    if v.endswith("%"):
        try:
            pct = float(v.replace("%", "").strip())
        except ValueError:
            pct = 0.0
        if pct <= 0:
            return "Low"
        if pct < 5:
            return "Medium"
        if pct < 10:
            return "High"
        return "High"

    if any(x in v for x in ["daily", "per day", "every day", "frequent", "often", "always"]):
        return "High"
    if any(x in v for x in ["weekly", "per week"]):
        return "High"
    if any(x in v for x in ["monthly", "per month"]):
        return "Medium"
    if any(x in v for x in ["quarterly", "annually", "annual", "per year"]):
        return "Low"

    try:
        n = int(float(re.sub(r"[^0-9.]", "", v)))
    except ValueError:
        return "Medium"

    if n <= 0:
        return "Low"
    if n == 1:
        return "Low"
    if 2 <= n <= 3:
        return "Medium"
    if 4 <= n <= 9:
        return "High"
    return "High"


def _impact_from_profile_text(text):
    t = (text or "").lower()
    very_high = [
        "money laundering", "aml", "cft", "terrorist financing", "sanction", "pep",
        "politically exposed", "regulatory penalty", "regulatory breach", "fraud",
        "identity theft", "data breach", "loss of funds", "unclear source of funds",
        "undocumented source", "complex ownership", "beneficial owner",
    ]
    high = [
        "high net worth", "high value", "kyc", "watchlist", "watch list", "reputational",
        "litigation", "legal", "regulatory", "customer due diligence", "enhanced due diligence",
        "transaction monitoring", "red-flag", "red flag", "mobile banking", "ussd",
    ]
    medium = [
        "documentation", "exception", "process", "delay", "control gap", "monitoring",
        "reporting", "review", "overdue", "complaint",
    ]
    if any(k in t for k in very_high):
        return "High"
    if any(k in t for k in high):
        return "High"
    if any(k in t for k in medium):
        return "Medium"
    return "Medium"


def _reduce_level(level):
    order = ["Low", "Medium", "High"]
    if level not in order:
        level = "Medium"
    return order[max(order.index(level) - 1, 0)]


def _extract_risk_records_from_text(raw_text, skip_zero_occurrence=False, source_type="", source_filename=""):
    profile_record = _parse_customer_risk_profile(raw_text, source_type=source_type, source_filename=source_filename)
    if profile_record:
        profile_record["duplicate_profile"] = _profile_duplicate_context(profile_record.get("customer_profile_data") or {})
        return {
            "area_name": profile_record["area_name"],
            "reporting_period": "",
            "results": [profile_record],
        }

    lines = [ln.strip() for ln in raw_text.splitlines() if ln.strip()]
    area_name = ""
    reporting_period = ""

    for ln in lines[:5]:
        if "Reporting Period:" in ln:
            left, right = ln.split("Reporting Period:", 1)
            area_name = left.strip()
            reporting_period = right.strip()
            break
    if not area_name and lines:
        area_name = lines[0].strip()

    header_idx = -1
    for i, ln in enumerate(lines):
        header_text = ln.lower()
        if "key risk indicator" in header_text or ("kri description" in header_text and "related risk" in header_text):
            header_idx = i
            break

    data_lines = lines[header_idx + 1:] if header_idx != -1 and header_idx + 1 < len(lines) else lines[1:]
    extracted = []
    counter = 1

    for ln in data_lines:
        row_text = ln.lower()
        if "key risk indicator" in row_text or ("kri description" in row_text and "related risk" in row_text):
            continue

        parts = _split_extracted_row(ln)
        if len(parts) < 3:
            continue

        kri = parts[0] if len(parts) >= 1 else ""
        kri_desc = parts[1] if len(parts) >= 2 else ""
        related_risk = parts[2] if len(parts) >= 3 else ""
        process = parts[3] if len(parts) >= 4 else ""
        occurrence = parts[4] if len(parts) >= 5 else ""
        combined = " ".join([kri, kri_desc, related_risk, process, occurrence])

        if skip_zero_occurrence and is_zero_occurrence(occurrence):
            continue

        base_ref = f"RISK-{_clean_ref_part(area_name)}-{counter:03d}"
        reference_id = make_unique_reference_id(base_ref)
        prob = _likelihood_from_occurrence(occurrence)
        impact = _impact_from_profile_text(combined)
        residual_prob = _reduce_level(prob) if not is_zero_occurrence(occurrence) else prob
        residual_impact = _reduce_level(impact) if impact == "High" else impact

        extracted.append({
            "reference_id": reference_id,
            "area_name": area_name,
            "reporting_period": reporting_period,
            "risk_owner": suggest_risk_owner(area_name),
            "risk_coordinator_name": "Risk & Compliance Coordinator",
            "risk_description": related_risk.strip() or kri.strip() or "TBD",
            "root_cause": kri_desc.strip(),
            "trigger": f"When {kri.strip().lower()} happens or is detected." if kri.strip() else "",
            "consequences": related_risk.strip(),
            "inherent_probability": prob,
            "inherent_impact": impact,
            "inherent_rating": "-",
            "control_descriptions": default_controls_for_area(area_name),
            "control_owner": suggest_risk_owner(area_name),
            "residual_probability": residual_prob,
            "residual_impact": residual_impact,
            "residual_rating": "-",
            "source_kri": kri,
            "source_kri_description": kri_desc,
            "source_related_risk": related_risk,
            "source_process": process,
            "source_occurrence": occurrence,
            "customer_profile_score": None,
            "customer_profile_rating": "",
            "customer_profile_notes": "",
            "customer_profile_data": None,
            "recommendation": "Review extracted risk and confirm controls, owner, and action plan before approval.",
            "confidence_notes": "KRI extraction used occurrence text, risk keywords, and process context to estimate likelihood and impact.",
        })
        counter += 1

    return {"area_name": area_name, "reporting_period": reporting_period, "results": extracted}


def _create_risk_from_extracted(result, request, workflow_status, evidence_file=None):
    risk = RiskAssessment.objects.create(
        reference_id=make_unique_reference_id(result["reference_id"]),
        area_name=result["area_name"],
        description=("[DRAFT] " if workflow_status == "Draft" else "") + result["risk_description"],
        caused_by=result["root_cause"],
        consequences=result["consequences"],
        customer_profile_score=result.get("customer_profile_score"),
        customer_profile_rating=result.get("customer_profile_rating", ""),
        customer_profile_notes=result.get("customer_profile_notes", ""),
        risk_owner=result["risk_owner"],
        risk_coordinator_name=result.get("risk_coordinator_name", ""),
        inherent_probability=result["inherent_probability"],
        inherent_impact=result["inherent_impact"],
        residual_probability=result["residual_probability"],
        residual_impact=result["residual_impact"],
        workflow_status=workflow_status,
        controls=result["control_descriptions"],
        control_owner=result["control_owner"],
        updated_by=request.user,
    )
    profile_data = result.get("customer_profile_data") or {}
    if profile_data:
        _create_customer_profile_from_data(profile_data, request, risk=risk, evidence_file=evidence_file)
    return risk


def _create_customer_profile_from_data(profile_data, request, risk=None, evidence_file=None):
    previous = _previous_customer_profile(
        account_no=profile_data.get("account_no", ""),
        account_name=profile_data.get("account_name", ""),
    )
    profile_rating = profile_data.get("profile_rating", "")
    movement = "New"
    if previous:
        movement = "Unchanged"
        if _profile_rank(profile_rating) > _profile_rank(previous.profile_rating):
            movement = "Worsened"
        elif _profile_rank(profile_rating) < _profile_rank(previous.profile_rating):
            movement = "Improved"
    review_frequency, next_review_date = _review_plan_for_rating(profile_rating)
    profile = CustomerRiskProfile.objects.create(
        risk=risk,
        account_no=profile_data.get("account_no", ""),
        account_name=profile_data.get("account_name", ""),
        profile_rating=profile_rating,
        average_score=profile_data.get("average_score"),
        total_score=profile_data.get("total_score") or 0,
        determinant_count=profile_data.get("determinant_count") or 0,
        determinants=profile_data.get("determinants") or [],
        recommendation=profile_data.get("recommendation", ""),
        enhanced_due_diligence=profile_data.get("enhanced_due_diligence", ""),
        confidence_notes=profile_data.get("confidence_notes", ""),
        source_filename=profile_data.get("source_filename", ""),
        source_type=profile_data.get("source_type", ""),
        evidence_file=evidence_file,
        review_frequency=review_frequency,
        next_review_date=next_review_date,
        previous_rating=previous.profile_rating if previous else "",
        previous_average_score=previous.average_score if previous else None,
        movement=movement,
        created_by=request.user,
    )
    log_system_event(
        request,
        "create",
        f"Saved customer risk profile for {profile.account_name or profile.account_no or 'customer'} ({profile.profile_rating}).",
        target=risk,
        metadata={"profile_id": profile.id, "movement": movement, "next_review_date": str(next_review_date)},
        area_name="Customer Risk Profile",
    )
    return profile


# --- LOGIN REDIRECT ---
def redirect_to_login(request):
    return redirect('/accounts/login/')


# --- DASHBOARD ---
@login_required
def dashboard(request):
    risks = RiskAssessment.objects.all().order_by('reference_id')


    selected_area = request.GET.get("area", "").strip()
    available_areas = list(
        RiskAssessment.objects.exclude(area_name__isnull=True)
        .exclude(area_name__exact="")
        .values_list("area_name", flat=True)
        .distinct()
    )
    if selected_area:
        risks = risks.filter(area_name=selected_area)

    filter_type = request.GET.get("filter", "all").strip()
    if filter_type == "draft":
        risks = risks.filter(workflow_status='Draft')
    elif filter_type == "approved":
        risks = risks.filter(workflow_status='Approved')
    elif filter_type == "reviewed":
        risks = risks.filter(workflow_status='Reviewed')
    elif filter_type == "closed":
        risks = risks.filter(workflow_status='Closed')

    risks = _apply_dashboard_filters(request, risks)

    probabilities = HEATMAP_PROBABILITIES
    impacts = HEATMAP_IMPACTS

    def get_matrix_counts(risk_type):
        matrix_grid = {p: {i: 0 for i in impacts} for p in probabilities}
        for r in risks:
            if risk_type == 'inherent':
                p, i = r.inherent_probability, r.inherent_impact
            else:
                p, i = r.residual_probability, r.residual_impact

            if p in matrix_grid and i in matrix_grid[p]:
                matrix_grid[p][i] += 1
        return matrix_grid

    def build_department_scores():
        scored_departments = []
        departments = (
            RiskAssessment.objects.exclude(area_name__isnull=True)
            .exclude(area_name__exact="")
            .values_list("area_name", flat=True)
            .distinct()
        )

        for area in departments:
            area_risks = list(RiskAssessment.objects.filter(area_name=area))
            total = len(area_risks)
            high = sum(1 for risk in area_risks if is_high_risk(risk.residual_rating))
            medium = sum(1 for risk in area_risks if normalize_risk_level(risk.residual_rating) == "Medium")
            overdue = sum(1 for risk in area_risks if risk.is_action_overdue)
            open_actions_count = sum(
                1 for risk in area_risks
                if risk.mitigation_action and risk.action_status != "Completed"
            )
            score = (high * 20) + (medium * 10) + (overdue * 10) + (open_actions_count * 3) + total

            scored_departments.append({
                "area": area,
                "score": score,
                "total": total,
                "high": high,
                "medium": medium,
                "overdue": overdue,
                "open_actions": open_actions_count,
            })

        return sorted(scored_departments, key=lambda item: (-item["score"], item["area"]))

    visible_risk_list = list(risks)
    incident_qs = RiskIncident.objects.all()
    if selected_area:
        incident_qs = incident_qs.filter(area_name=selected_area)
    total_loss_amount = sum(incident.loss_amount for incident in incident_qs)
    early_warnings = {
        "high": [risk for risk in visible_risk_list if is_high_risk(risk.residual_rating)][:10],
        "overdue": [risk for risk in visible_risk_list if risk.is_action_overdue][:10],
        "worsening": [risk for risk in visible_risk_list if risk.residual_trend_label == "Worsening"][:10],
        "missing_owner": [
            risk for risk in visible_risk_list
            if not (risk.risk_owner or "").strip() or not (risk.control_owner or "").strip()
        ][:10],
        "missing_action": [
            risk for risk in visible_risk_list
            if is_medium_or_high_risk(risk.residual_rating) and not (risk.mitigation_action or "").strip()
        ][:10],
        "weak_controls": [
            risk for risk in visible_risk_list
            if risk.control_effectiveness == "Weak"
        ][:10],
        "escalated": [
            risk for risk in visible_risk_list
            if risk.escalation_status != "Normal"
        ][:10],
    }
    notifications = []
    for risk in early_warnings["overdue"][:5]:
        notifications.append({"level": "warning", "title": f"{risk.reference_id} overdue", "body": f"{risk.days_overdue} day(s) overdue"})
    for risk in early_warnings["high"][:5]:
        notifications.append({"level": "danger", "title": f"{risk.reference_id} high risk", "body": risk.area_name or "Unspecified"})
    for incident in incident_qs.exclude(status__in=["Resolved", "Closed"])[:5]:
        notifications.append({"level": "info", "title": f"Incident: {incident.title}", "body": incident.area_name or "Unspecified"})
    for risk in risks.filter(workflow_status__in=["Draft", "Reviewed"])[:5]:
        notifications.append({"level": "info", "title": f"{risk.reference_id} pending {risk.workflow_status}", "body": risk.area_name or "Unspecified"})

    department_chart = build_department_scores()[:8]
    overdue_trend = [
        {"label": "Overdue", "count": sum(1 for risk in visible_risk_list if risk.is_action_overdue)},
        {"label": "Next 30 Days", "count": sum(1 for risk in visible_risk_list if risk.action_due_date and timezone.localdate() <= risk.action_due_date <= timezone.localdate() + timedelta(days=30) and risk.action_status != "Completed")},
        {"label": "No Due Date", "count": sum(1 for risk in visible_risk_list if not risk.action_due_date and risk.action_status != "Completed")},
    ]
    incident_trend = [
        {"label": incident.incident_date.strftime("%Y-%m-%d"), "count": 1, "loss": incident.loss_amount}
        for incident in incident_qs.order_by("-incident_date", "-created_at")[:8]
    ]

    assurance_items = _build_assurance_cockpit(visible_risk_list, incident_qs, available_areas)

    context = {
        'risks': risks,
        'total_risks': risks.count(),
        'high_risks': sum(1 for risk in risks if is_high_risk(risk.residual_rating)),
        'overdue_actions': sum(1 for risk in risks if risk.is_action_overdue),
        'escalated_risks': sum(1 for risk in risks if risk.escalation_status != "Normal"),
        'open_actions': risks.exclude(action_status='Completed').exclude(mitigation_action__exact='').count(),
        'total_incidents': incident_qs.count(),
        'open_incidents': incident_qs.exclude(status__in=['Resolved', 'Closed']).count(),
        'total_loss_amount': total_loss_amount,
        'user': request.user,
        'probabilities': probabilities,
        'impacts': impacts,
        'inherent_matrix': get_matrix_counts('inherent'),
        'residual_matrix': get_matrix_counts('residual'),
        'department_scores': build_department_scores(),
        'early_warnings': early_warnings,
        'available_areas': available_areas,
        'selected_area': selected_area,
        'filter_type': filter_type,
        'rating_chart': _rating_chart(visible_risk_list),
        'department_chart': department_chart,
        'overdue_trend': overdue_trend,
        'incident_trend': incident_trend,
        'notifications': notifications[:12],
        'assurance_items': assurance_items,
        'notification_count': len(notifications),
        'query_text': request.GET.get("q", "").strip(),
        'selected_owner': request.GET.get("owner", "").strip(),
        'selected_rating': request.GET.get("rating", "").strip(),
        'selected_action_status': request.GET.get("action_status", "").strip(),
        'selected_escalation': request.GET.get("escalation", "").strip(),
        'selected_due': request.GET.get("due", "").strip(),
        'rating_choices': RiskAssessment.RATING_CHOICES,
        'action_status_choices': RiskAssessment.ACTION_STATUS_CHOICES,
        'escalation_status_choices': RiskAssessment.ESCALATION_STATUS_CHOICES,
        'incident_status_choices': RiskIncident.INCIDENT_STATUS_CHOICES,
        'can_submit_risk': can_submit_risk(request.user),
        'can_review_risk': can_review_risk(request.user),
        'can_approve_delete': can_approve_delete(request.user),
    }
    return render(request, 'risks/dashboard.html', context)


@login_required
def compliance_assistant(request):
    from .copilot import answer_compliance_question, build_snapshot

    suggested_questions = [
        "Explain the highest risks and why management should care.",
        "Which risks need urgent attention today?",
        "Use internet context to suggest controls for the high risks.",
        "Compare our risk exposure with common banking risk practices.",
    ]
    question = ""
    answer = "Ask the AI about risks, departments, controls, incidents, actions, reports, or current external risk context."
    assistant_meta = {
        "used_live_llm": False,
        "provider": "system",
        "model_name": "grounded",
        "response_ms": 0,
        "error_message": "",
        "citations": [],
    }

    if request.method == "POST":
        question = request.POST.get("question", "").strip()
        result = answer_compliance_question(user=request.user, question=question)
        answer = result["answer"]
        assistant_meta = {
            "used_live_llm": result["used_live_llm"],
            "provider": result["provider"],
            "model_name": result["model_name"],
            "response_ms": result["response_ms"],
            "error_message": result["error_message"],
            "citations": result["citations"],
        }

    return render(request, "risks/compliance_assistant.html", {
        "snapshot": build_snapshot(),
        "suggested_questions": suggested_questions,
        "question": question,
        "answer": answer,
        "assistant_meta": assistant_meta,
    })


@login_required
def risk_create(request):
    if not can_submit_risk(request.user):
        return HttpResponseForbidden("<h1>Access Denied</h1>")
    if request.method == "POST":
        form = RiskAssessmentForm(request.POST)
        if form.is_valid():
            risk = form.save(commit=False)
            risk.updated_by = request.user
            if not can_review_risk(request.user):
                risk.workflow_status = "Draft"
            risk.save()
            log_system_event(request, "create", f"Created risk {risk.reference_id}.", target=risk)
            messages.success(request, f"Risk {risk.reference_id} was saved.")
            return redirect("risk-detail", risk_id=risk.id)
    else:
        form = RiskAssessmentForm(initial={"workflow_status": "Draft"})
    return render(request, "risks/risk_form.html", {"form": form, "title": "Add Risk", "mode": "create"})


@login_required
def risk_detail(request, risk_id):
    risk = get_object_or_404(RiskAssessment, id=risk_id)
    return render(request, "risks/risk_detail.html", {
        "risk": risk,
        "incidents": risk.incidents.all(),
        "customer_profiles": risk.customer_profiles.all()[:10],
        "audit_logs": SystemAuditLog.objects.filter(reference_id=risk.reference_id)[:15],
        "can_submit_risk": can_submit_risk(request.user),
        "can_review_risk": can_review_risk(request.user),
        "can_approve_delete": can_approve_delete(request.user),
    })


@login_required
def customer_profile_list(request):
    q = request.GET.get("q", "").strip()
    profiles = CustomerRiskProfile.objects.select_related("risk", "created_by").all()
    if q:
        profiles = profiles.filter(
            Q(account_name__icontains=q)
            | Q(account_no__icontains=q)
            | Q(profile_rating__icontains=q)
            | Q(risk__reference_id__icontains=q)
        )
    return render(request, "risks/customer_profiles.html", {"profiles": profiles[:250], "query_text": q})


@login_required
def customer_profile_detail(request, profile_id):
    profile = get_object_or_404(
        CustomerRiskProfile.objects.select_related("risk", "created_by"),
        id=profile_id,
    )
    history = CustomerRiskProfile.objects.filter(
        Q(account_no=profile.account_no) | Q(account_name=profile.account_name)
    ).exclude(id=profile.id)[:20]
    return render(request, "risks/customer_profile_detail.html", {"profile": profile, "history": history})


@login_required
def customer_profile_report(request, profile_id):
    profile = get_object_or_404(
        CustomerRiskProfile.objects.select_related("risk", "created_by"),
        id=profile_id,
    )
    history = CustomerRiskProfile.objects.filter(
        Q(account_no=profile.account_no) | Q(account_name=profile.account_name)
    ).exclude(id=profile.id)[:10]
    return render(request, "risks/customer_profile_report.html", {
        "profile": profile,
        "history": history,
        "generated_at": timezone.now(),
        "generated_by": request.user,
    })


@login_required
def risk_update(request, risk_id):
    if not can_review_risk(request.user):
        return HttpResponseForbidden("<h1>Access Denied</h1>")
    risk = get_object_or_404(RiskAssessment, id=risk_id)
    if request.method == "POST":
        form = RiskAssessmentForm(request.POST, instance=risk)
        if form.is_valid():
            risk = form.save(commit=False)
            risk.updated_by = request.user
            risk.save()
            log_system_event(request, "update", f"Updated risk {risk.reference_id}.", target=risk)
            messages.success(request, f"Risk {risk.reference_id} was updated.")
            return redirect("risk-detail", risk_id=risk.id)
    else:
        form = RiskAssessmentForm(instance=risk)
    return render(request, "risks/risk_form.html", {"form": form, "risk": risk, "title": f"Edit {risk.reference_id}", "mode": "edit"})


@login_required
def risk_action_update(request, risk_id):
    if not can_review_risk(request.user):
        return HttpResponseForbidden("<h1>Access Denied</h1>")
    risk = get_object_or_404(RiskAssessment, id=risk_id)
    if request.method == "POST":
        form = RiskActionUpdateForm(request.POST, instance=risk)
        if form.is_valid():
            risk = form.save(commit=False)
            risk.updated_by = request.user
            risk.save()
            log_system_event(request, "update", f"Updated action plan for {risk.reference_id}.", target=risk)
            messages.success(request, "Action update saved.")
            return redirect("risk-detail", risk_id=risk.id)
    else:
        form = RiskActionUpdateForm(instance=risk)
    return render(request, "risks/risk_form.html", {"form": form, "risk": risk, "title": f"Action Update - {risk.reference_id}", "mode": "action"})


@login_required
def incident_create(request, risk_id=None):
    if not can_submit_risk(request.user):
        return HttpResponseForbidden("<h1>Access Denied</h1>")
    risk = get_object_or_404(RiskAssessment, id=risk_id) if risk_id else None
    if request.method == "POST":
        form = RiskIncidentForm(request.POST)
        if form.is_valid():
            incident = form.save(commit=False)
            if risk:
                incident.risk = risk
                incident.area_name = incident.area_name or risk.area_name or ""
            incident.save()
            log_system_event(request, "create", f"Recorded incident {incident.title}.", target=incident, area_name=incident.area_name)
            messages.success(request, "Incident recorded.")
            return redirect("risk-detail", risk_id=risk.id) if risk else redirect("risk-calendar")
    else:
        form = RiskIncidentForm(initial={"risk": risk, "area_name": risk.area_name if risk else "", "incident_date": timezone.localdate()})
    return render(request, "risks/incident_form.html", {"form": form, "risk": risk})


@login_required
def risk_delete(request, risk_id):
    if not can_approve_delete(request.user):
        return HttpResponseForbidden("<h1>Access Denied</h1>")
    risk = get_object_or_404(RiskAssessment, id=risk_id)
    if request.method == "POST":
        confirm_text = request.POST.get("confirm_text", "").strip()
        delete_reason = request.POST.get("delete_reason", "").strip()
        if confirm_text != "DELETE" or not delete_reason:
            messages.error(request, "Type DELETE and provide a reason.")
            return redirect("risk-delete", risk_id=risk.id)
        reference_id = risk.reference_id
        area_name = risk.area_name or ""
        risk.delete()
        log_system_event(request, "delete", f"Deleted risk {reference_id}.", metadata={"reason": delete_reason}, area_name=area_name)
        messages.warning(request, f"Risk {reference_id} was deleted.")
        return redirect("dashboard")
    return render(request, "risks/risk_confirm_delete.html", {"risk": risk})


@login_required
def department_detail(request, area_name):
    risks = RiskAssessment.objects.filter(area_name=area_name).order_by("reference_id")
    incidents = RiskIncident.objects.filter(area_name=area_name)
    risk_list = list(risks)
    board_summary = (
        f"{area_name} has {len(risk_list)} recorded risk item(s), "
        f"{sum(1 for risk in risk_list if is_high_risk(risk.residual_rating))} high residual risk(s), "
        f"{sum(1 for risk in risk_list if risk.is_action_overdue)} overdue action(s), and "
        f"{incidents.exclude(status__in=['Resolved', 'Closed']).count()} open incident(s)."
    )
    probabilities = HEATMAP_PROBABILITIES
    impacts = HEATMAP_IMPACTS
    matrix = {p: {i: 0 for i in impacts} for p in probabilities}
    for risk in risk_list:
        if risk.residual_probability in matrix and risk.residual_impact in matrix[risk.residual_probability]:
            matrix[risk.residual_probability][risk.residual_impact] += 1
    return render(request, "risks/department_detail.html", {
        "area_name": area_name,
        "risks": risks,
        "incidents": incidents,
        "score": _department_score(area_name, risk_list),
        "rating_chart": _rating_chart(risk_list),
        "overdue_actions": [risk for risk in risk_list if risk.is_action_overdue],
        "board_summary": board_summary,
        "probabilities": probabilities,
        "impacts": impacts,
        "residual_matrix": matrix,
    })


@login_required
def risk_calendar(request):
    risks = RiskAssessment.objects.exclude(action_due_date__isnull=True).order_by("action_due_date", "reference_id")
    incidents = RiskIncident.objects.all().order_by("-incident_date", "-created_at")[:50]
    board_deadlines = [
        {"date": timezone.localdate(), "title": "Today"},
        {"date": timezone.localdate() + timedelta(days=30), "title": "Next board reporting checkpoint"},
    ]
    return render(request, "risks/risk_calendar.html", {"risks": risks, "incidents": incidents, "board_deadlines": board_deadlines, "today": timezone.localdate()})


@login_required
def audit_log(request):
    if not can_review_risk(request.user) and not _in_group(request.user, "Board"):
        return HttpResponseForbidden("<h1>Access Denied</h1>")
    logs = SystemAuditLog.objects.all()
    q = request.GET.get("q", "").strip()
    action = request.GET.get("action", "").strip()
    if q:
        logs = logs.filter(Q(summary__icontains=q) | Q(reference_id__icontains=q) | Q(area_name__icontains=q) | Q(user__username__icontains=q))
    if action:
        logs = logs.filter(action=action)
    return render(request, "risks/audit_log.html", {"logs": logs[:250], "query_text": q, "selected_action": action, "action_choices": SystemAuditLog.ACTION_CHOICES})


@login_required
def notifications_panel(request):
    risks = list(RiskAssessment.objects.all())
    incidents = RiskIncident.objects.exclude(status__in=["Resolved", "Closed"])[:20]
    notifications = {
        "overdue": [risk for risk in risks if risk.is_action_overdue],
        "high": [risk for risk in risks if is_high_risk(risk.residual_rating)],
        "incidents": incidents,
        "pending": [risk for risk in risks if risk.workflow_status in ["Draft", "Reviewed"]],
    }
    return render(request, "risks/notifications.html", {"notifications": notifications})


@login_required
def backup_tools(request):
    if not can_approve_delete(request.user):
        return HttpResponseForbidden("<h1>Access Denied</h1>")
    return render(request, "risks/backup_tools.html")


@login_required
def backup_database(request):
    if not can_approve_delete(request.user):
        return HttpResponseForbidden("<h1>Access Denied</h1>")
    db_path = settings.DATABASES["default"].get("NAME")
    log_system_event(request, "export", "Downloaded SQLite database backup.")
    return FileResponse(open(db_path, "rb"), as_attachment=True, filename="bank_risk_system_backup.sqlite3")


@login_required
def restore_database(request):
    if not can_approve_delete(request.user):
        return HttpResponseForbidden("<h1>Access Denied</h1>")
    if request.method != "POST":
        return redirect("backup-tools")
    upload = request.FILES.get("database_file")
    confirm_text = request.POST.get("confirm_text", "").strip()
    reason = request.POST.get("restore_reason", "").strip()
    if not upload or confirm_text != "RESTORE" or not reason:
        messages.error(request, "Upload a database file, give a reason, and type RESTORE.")
        return redirect("backup-tools")
    db_path = settings.DATABASES["default"].get("NAME")
    safety_backup = f"{db_path}.{timezone.now().strftime('%Y%m%d%H%M%S')}.bak"
    shutil.copy2(db_path, safety_backup)
    with open(db_path, "wb") as destination:
        for chunk in upload.chunks():
            destination.write(chunk)
    log_system_event(request, "update", "Restored SQLite database backup.", metadata={"reason": reason, "safety_backup": safety_backup})
    messages.warning(request, "Database restored. Restart the server if needed.")
    return redirect("dashboard")


@login_required
def export_system_package(request):
    if not can_approve_delete(request.user):
        return HttpResponseForbidden("<h1>Access Denied</h1>")
    base_dir = settings.BASE_DIR
    temp = tempfile.NamedTemporaryFile(delete=False, suffix=".zip")
    temp.close()
    skip_dirs = {"__pycache__", ".git", "staticfiles"}
    with zipfile.ZipFile(temp.name, "w", zipfile.ZIP_DEFLATED) as archive:
        for root, dirs, files in os.walk(base_dir):
            dirs[:] = [d for d in dirs if d not in skip_dirs]
            for file_name in files:
                file_path = os.path.join(root, file_name)
                if file_path == temp.name:
                    continue
                archive.write(file_path, os.path.relpath(file_path, base_dir))
    log_system_event(request, "export", "Exported full system package.")
    return FileResponse(open(temp.name, "rb"), as_attachment=True, filename="bank_risk_system_package.zip")


# --- EXPORT CSV ---
@login_required
def export_risks_csv(request):
    selected_area = request.GET.get("area", "").strip()
    export_qs = RiskAssessment.objects.all().order_by('-created_at')
    if selected_area:
        export_qs = export_qs.filter(area_name=selected_area)

    log_system_event(
        request,
        "export",
        f"Exported {export_qs.count()} risk record(s).",
        metadata={"area": selected_area or "All Departments"},
        area_name=selected_area,
    )

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="risk_register.csv"'
    writer = csv.writer(response)
    writer.writerow(['ID', 'Risk Category', 'Risk Description', 'Likelihood', 'Impact', 'Risk Rank', 'Risk Trigger', 'Preventive Controls', 'Risk Treatment', 'Risk Owner', 'Risk Coordinator', 'Residual Risk'])

    for risk in export_qs:
        writer.writerow([
            risk.reference_id,
            risk.area_name,
            risk.description,
            display_level(risk.inherent_probability),
            display_level(risk.inherent_impact),
            display_level(risk.inherent_rating),
            risk.caused_by,
            risk.control_description,
            risk.mitigation_action or "mitigate",
            risk.risk_owner,
            risk.risk_coordinator_name or "Risk & Compliance Officer",
            display_level(risk.residual_rating)
        ])
    return response


@login_required
def executive_dashboard(request):
    risks = list(RiskAssessment.objects.all().order_by("area_name", "reference_id"))
    incidents = list(RiskIncident.objects.all())
    ratings = RISK_LEVELS

    rating_counts = {
        rating: sum(1 for risk in risks if risk.residual_rating == rating)
        for rating in ratings
    }

    department_rows = []
    departments = sorted({risk.area_name or "Unspecified" for risk in risks})
    for area in departments:
        area_risks = [risk for risk in risks if (risk.area_name or "Unspecified") == area]
        high = sum(1 for risk in area_risks if is_high_risk(risk.residual_rating))
        medium = sum(1 for risk in area_risks if normalize_risk_level(risk.residual_rating) == "Medium")
        overdue = sum(1 for risk in area_risks if risk.is_action_overdue)
        escalated = sum(1 for risk in area_risks if risk.escalation_status != "Normal")
        score = (high * 20) + (medium * 10) + (overdue * 10) + (escalated * 10) + len(area_risks)
        department_rows.append({
            "area": area,
            "score": score,
            "total": len(area_risks),
            "high": high,
            "medium": medium,
            "overdue": overdue,
            "escalated": escalated,
        })

    department_rows.sort(key=lambda item: (-item["score"], item["area"]))

    context = {
        "total_risks": len(risks),
        "rating_counts": rating_counts,
        "department_rows": department_rows,
        "top_risks": sorted(
            risks,
            key=lambda risk: (
                {"High": 0, "Medium": 1, "Low": 2}.get(normalize_risk_level(risk.residual_rating), 9),
                risk.reference_id,
            ),
        )[:10],
        "overdue_actions": [risk for risk in risks if risk.is_action_overdue][:10],
        "escalated_risks": [risk for risk in risks if risk.escalation_status != "Normal"][:10],
        "incident_count": len(incidents),
        "open_incident_count": sum(1 for incident in incidents if incident.status not in ["Resolved", "Closed"]),
        "total_loss_amount": sum(incident.loss_amount for incident in incidents),
    }
    return render(request, "risks/executive_dashboard.html", context)


# --- OFFICIAL REPORT ---
@login_required
def official_report(request):
    if not request.user.is_superuser and not request.user.has_perm('risks.view_reportconfiguration'):
        return HttpResponseForbidden("<h1>Access Denied</h1><p>You do not have permission to view this official document.</p>")

    config, created = ReportConfiguration.objects.get_or_create(id=1)

    if request.method == "POST" and request.user.is_superuser:
        new_summary = request.POST.get('executive_summary')
        if new_summary:
            config.executive_summary = new_summary
            config.save()

    risks = RiskAssessment.objects.all().order_by('area_name', 'reference_id')
    incidents = RiskIncident.objects.all()


    # group by area_name for headings
    grouped = {}
    for r in risks:
        key = r.area_name or "UNSPECIFIED"
        grouped.setdefault(key, []).append(r)

    visible_risks = list(risks)
    report_quality_warnings = [
        {
            "label": "Missing owners",
            "count": sum(1 for risk in visible_risks if not (risk.risk_owner or "").strip()),
        },
        {
            "label": "Missing controls",
            "count": sum(1 for risk in visible_risks if not (risk.controls or "").strip()),
        },
        {
            "label": "No action plan",
            "count": sum(
                1 for risk in visible_risks
                if is_medium_or_high_risk(risk.residual_rating) and not (risk.mitigation_action or "").strip()
            ),
        },
        {
            "label": "Overdue actions",
            "count": sum(1 for risk in visible_risks if risk.is_action_overdue),
        },
        {
            "label": "Stale reviews",
            "count": sum(1 for risk in visible_risks if risk.updated_at.date() < timezone.localdate() - timedelta(days=90)),
        },
    ]

    context = {
        'risks': risks,
        'grouped_risks': grouped,
        'config': config,
        'generated_at': timezone.now(),
        'generated_by': request.user.username,
        'is_admin': request.user.is_superuser,
        'report_quality_warnings': report_quality_warnings,
        'report_summary': {
            'total_risks': risks.count(),
            'high_risks': sum(1 for risk in risks if is_high_risk(risk.residual_rating)),
            'medium_risks': sum(1 for risk in risks if normalize_risk_level(risk.residual_rating) == 'Medium'),
            'low_risks': sum(1 for risk in risks if normalize_risk_level(risk.residual_rating) == 'Low'),
            'overdue_actions': sum(1 for risk in risks if risk.is_action_overdue),
            'escalated_risks': sum(1 for risk in risks if risk.escalation_status != 'Normal'),
            'incident_count': incidents.count(),
            'open_incident_count': incidents.exclude(status__in=['Resolved', 'Closed']).count(),
            'total_loss_amount': sum(incident.loss_amount for incident in incidents),
        },
    }
    return render(request, 'admin/official_report.html', context)


# ========= AI EXTRACT (Preview) =========
@login_required
def ai_extract_risks(request):
    context = {
        "raw_text": "",
        "area_name": "",
        "reporting_period": "",
        "results": [],
        "error": "",
        "source_filename": "",
        "source_type": "",
    }

    if request.method == "POST":
        raw_text = request.POST.get("raw_text", "")
        uploaded_file = request.FILES.get("source_file")
        source_filename = ""
        source_type = "pasted text" if raw_text.strip() else ""
        if uploaded_file:
            try:
                file_text, source_filename, source_type = _uploaded_file_to_text(uploaded_file)
                raw_text = f"{raw_text}\n{file_text}".strip() if raw_text.strip() else file_text
            except Exception as exc:
                context["error"] = f"Could not read uploaded file: {exc}"

        context["raw_text"] = raw_text
        context["source_filename"] = source_filename
        context["source_type"] = source_type

        if context["error"]:
            pass
        elif not raw_text.strip():
            context["error"] = "Please paste KRI/profile text or upload an Excel, Word, PDF, or TXT file first."
        else:
            extraction = _extract_risk_records_from_text(raw_text, source_type=source_type, source_filename=source_filename)
            context["area_name"] = extraction["area_name"]
            context["reporting_period"] = extraction["reporting_period"]
            context["results"] = extraction["results"]
            if not extraction["results"]:
                context["error"] = "I could not detect any table rows. Make sure you pasted the KRI table with rows."

    return render(request, "risks/ai_extract.html", context)


# ========= SAVE DRAFTS =========
@login_required
def ai_extract_save_drafts(request):
    if request.method != "POST":
        return redirect("ai-extract")

    raw_text = request.POST.get("raw_text", "").strip()
    if not raw_text:
        return redirect("ai-extract")

    extraction = _extract_risk_records_from_text(
        raw_text,
        skip_zero_occurrence=True,
        source_type=request.POST.get("source_type", ""),
        source_filename=request.POST.get("source_filename", ""),
    )
    saved_count = 0
    evidence_file = request.FILES.get("evidence_file")
    for result in extraction["results"]:
        try:
            _create_risk_from_extracted(result, request, "Draft", evidence_file=evidence_file)
            saved_count += 1
        except Exception:
            continue

    return redirect(f"{reverse('dashboard')}?saved={saved_count}")


@login_required
def ai_extract_save_profile_only(request):
    if request.method != "POST":
        return redirect("ai-extract")

    raw_text = request.POST.get("raw_text", "").strip()
    evidence_file = request.FILES.get("evidence_file")
    source_filename = request.POST.get("source_filename", "")
    source_type = request.POST.get("source_type", "")
    if evidence_file and not raw_text:
        try:
            file_text, source_filename, source_type = _uploaded_file_to_text(evidence_file)
            evidence_file.seek(0)
            raw_text = file_text
        except Exception:
            return redirect("ai-extract")
    if not raw_text:
        return redirect("ai-extract")

    extraction = _extract_risk_records_from_text(
        raw_text,
        skip_zero_occurrence=False,
        source_type=source_type,
        source_filename=source_filename,
    )
    profile_result = next((item for item in extraction["results"] if item.get("customer_profile_data")), None)
    if not profile_result:
        return redirect("ai-extract")

    profile = _create_customer_profile_from_data(
        profile_result["customer_profile_data"],
        request,
        risk=None,
        evidence_file=evidence_file,
    )
    return redirect("customer-profile-detail", profile_id=profile.id)


# ========= SAVE & APPROVE =========
@login_required
def ai_extract_save_and_approve(request):
    
    if request.method != "POST":
        return redirect("ai-extract")

    raw_text = request.POST.get("raw_text", "")
    if not raw_text or not raw_text.strip():
        return redirect("ai-extract")

    extraction = _extract_risk_records_from_text(
        raw_text,
        skip_zero_occurrence=True,
        source_type=request.POST.get("source_type", ""),
        source_filename=request.POST.get("source_filename", ""),
    )
    saved_count = 0
    evidence_file = request.FILES.get("evidence_file")
    for result in extraction["results"]:
        try:
            _create_risk_from_extracted(result, request, "Approved", evidence_file=evidence_file)
            saved_count += 1
        except Exception:
            continue

    return redirect(f"{reverse('dashboard')}?saved={saved_count}")

    import re

    def split_row(line):
        if "\t" in line:
            return [p.strip() for p in line.split("\t") if p.strip()]
        return [p.strip() for p in re.split(r"\s{2,}", line.strip()) if p.strip()]

    def make_unique_reference_id(base_ref):
        ref = base_ref
        bump = 1
        while RiskAssessment.objects.filter(reference_id=ref).exists():
            ref = f"{base_ref}-{bump}"
            bump += 1
        return ref

    def likelihood_from_occurrence(value):
        v = str(value).strip().lower()

        # percentage like 10%
        if v.endswith("%"):
            try:
                pct = float(v.replace("%", "").strip())
            except ValueError:
                pct = 0.0
            if pct <= 0:
                return "Low"
            if pct < 5:
                return "Medium"
            if pct < 10:
                return "High"
            return "High"

        # frequency phrases
        if any(x in v for x in ["daily", "per day", "every day"]):
            return "High"
        if any(x in v for x in ["weekly", "per week", "frequently", "often"]):
            return "High"
        if any(x in v for x in ["monthly", "per month"]):
            return "Medium"
        if any(x in v for x in ["quarterly", "per quarter"]):
            return "Low"
        if any(x in v for x in ["annually", "annual", "per year"]):
            return "Low"

        # numeric
        try:
            n = int(v)
        except ValueError:
            # blank/unknown text -> Medium is safer than Low
            return "Medium"

        if n <= 0:
            return "Low"
        if n == 1:
            return "Low"
        if 2 <= n <= 3:
            return "Medium"
        if 4 <= n <= 9:
            return "High"
        return "High"

    def impact_from_text(text):
        t = (text or "").lower()

        very_high = [
            "money laundering", "aml", "cft", "sanction", "regulatory", "penalty",
            "fraud", "theft", "misappropriation", "terrorist financing",
            "data breach", "privacy breach", "identity theft", "loss of funds"
        ]
        high = [
            "legal", "contract", "reputational", "litigation", "complaint to the regulator",
            "regulatory scrutiny", "enforcement"
        ]
        medium = [
            "operational", "process", "delay", "reporting", "documentation", "control breakdown",
            "governance", "recommendation", "overdue corrective"
        ]

        if any(k in t for k in very_high):
            return "High"
        if any(k in t for k in high):
            return "High"
        if any(k in t for k in medium):
            return "Medium"
        if any(k in t for k in ["vault", "insurance", "cash exposure", "cash vault"]):
            return "High"

        return "Medium"

    def reduce_level(level):
        order = ["Low", "Medium", "High"]
        if level not in order:
            level = "Medium"
        return order[max(order.index(level) - 1, 0)]

    # ---------- PARSE LINES ----------
    lines = [ln.strip() for ln in raw_text.splitlines() if ln.strip()]
    if not lines:
        return redirect("ai-extract")

    # Parse area name safely from first line
    first = lines[0]
    if "Reporting Period:" in first:
        area_name = first.split("Reporting Period:", 1)[0].strip()
    else:
        area_name = first.strip()

    # Find header safely (no StopIteration)
    header_idx = -1
    for i, ln in enumerate(lines):
        if "Key Risk Indicator" in ln:
            header_idx = i
            break

    data_lines = lines[header_idx + 1:] if header_idx != -1 else lines[1:]

    OWNER_MAP = {
        "COMPLIANCE": "Compliance Manager",
        "AML": "Compliance Manager",
        "AUDIT": "Internal Auditor",
        "CREDIT": "Head of Credit",
        "LOAN RECOVERY": "Head of Credit",
        "SUSU": "Head of Operations",
        "OPERATIONAL": "Head of Operations",
        "IT": "Head of IT",
        "FINANCE": "Head of Finance",
        "TREASURY": "Head of Treasury",
    }

    # ========= COORDINATOR_MAP_START =========
    COORDINATOR_MAP = {
        # Compliance / AML
        "aml": "Compliance Officer",
        "cft": "Compliance Officer",
        "money laundering": "Compliance Officer",
        "sanction": "Compliance Officer",
        "regulatory": "Compliance Officer",
        "fic": "Compliance Officer",
        "bog": "Compliance Officer",

        # Fraud / theft
        "fraud": "Fraud & Investigations Officer",
        "theft": "Fraud & Investigations Officer",
        "misappropriation": "Fraud & Investigations Officer",
        "robbery": "Security Coordinator",

        # IT / systems
        "system": "IT Support Lead",
        "downtime": "IT Support Lead",
        "alert": "IT Support Lead",
        "verification system": "IT Support Lead",

        # Treasury / liquidity
        "liquidity": "Treasury Coordinator",
        "reserve": "Treasury Coordinator",
        "clearing": "Treasury Coordinator",
        "settlement": "Treasury Coordinator",

        # Customer / service
        "complaint": "Customer Service Coordinator",
        "reputational": "Customer Service Coordinator",

        # HR / people
        "staff": "HR Coordinator",
        "training": "HR Coordinator",
        "competency": "HR Coordinator",

        "__default__": "Risk & Compliance Coordinator",
    }
    # ========= COORDINATOR_MAP_END =========

    counter = 1

    for ln in data_lines:
        # ===== SKIP TABLE HEADER ROW =====
        if "kri description" in ln.lower() and "related risk" in ln.lower():
            continue
        # ================================

        parts = split_row(ln)
        if len(parts) < 3:
            continue

        kri = parts[0] if len(parts) >= 1 else ""
        kri_desc = parts[1] if len(parts) >= 2 else ""
        related_risk = parts[2] if len(parts) >= 3 else ""
        process = parts[3] if len(parts) >= 4 else ""
        occ = parts[4] if len(parts) >= 5 else ""

        # ========= OWNER_SELECT_START =========
        owner = "Department Head"
        for k, v in OWNER_MAP.items():
            if k in area_name.upper():
                owner = v
                break
        # ========= OWNER_SELECT_END =========

        # ========= COORDINATOR_SELECT_START =========
        combined_text = f"{kri} {kri_desc} {related_risk} {process}".lower()

        coordinator = COORDINATOR_MAP.get("__default__", "Risk Coordinator")
        for key, coord_name in COORDINATOR_MAP.items():
            if key != "__default__" and key in combined_text:
                coordinator = coord_name
                break
        # ========= COORDINATOR_SELECT_END =========

        # (then continue with your skip-zero check, scoring, and create())

        # ===== SKIP ZERO OCCURRENCE RISKS =====
        if is_zero_occurrence(occ):
            continue
        # =====================================


        inherent_prob = likelihood_from_occurrence(occ)
        inherent_impact = impact_from_text(" ".join([related_risk, kri, kri_desc, process]))

        residual_prob = reduce_level(inherent_prob)
        residual_impact = reduce_level(inherent_impact)

        base_ref = f"RISK-{area_name[:12].upper().replace(' ', '-')}-{counter:03d}"
        base_ref = re.sub(r"[^A-Z0-9\-]", "", base_ref)
        reference_id = make_unique_reference_id(base_ref)

        # Safe create: never crash whole request
        try:
            RiskAssessment.objects.create(
                reference_id=reference_id,
                area_name=area_name,
                description=related_risk or kri or "TBD",
                caused_by=kri_desc,
                consequences=related_risk,
                risk_owner=owner,
                risk_coordinator_name=coordinator,   # ✅ HERE
                inherent_probability=inherent_prob,
                inherent_impact=inherent_impact,
                residual_probability=residual_prob,
                residual_impact=residual_impact,
                workflow_status='Approved',
                controls="Standard Controls",
                control_owner=owner,
                updated_by=request.user
            )
        except Exception:
            pass

        counter += 1

    return redirect("dashboard")




@login_required
def edit_draft_risk(request, risk_id):
    risk = get_object_or_404(RiskAssessment, id=risk_id)

    if request.method == "POST":
        risk.area_name = request.POST.get("area_name", risk.area_name)
        risk.description = request.POST.get("description", risk.description)
        risk.caused_by = request.POST.get("caused_by", risk.caused_by)
        risk.consequences = request.POST.get("consequences", risk.consequences)
        risk.risk_owner = request.POST.get("risk_owner", risk.risk_owner)
        risk.controls = request.POST.get("controls", risk.controls)
        risk.control_owner = request.POST.get("control_owner", risk.control_owner)

        if request.POST.get("inherent_probability"):
            risk.inherent_probability = request.POST.get("inherent_probability")
        if request.POST.get("inherent_impact"):
            risk.inherent_impact = request.POST.get("inherent_impact")
        if request.POST.get("residual_probability"):
            risk.residual_probability = request.POST.get("residual_probability")
        if request.POST.get("residual_impact"):
            risk.residual_impact = request.POST.get("residual_impact")

        risk.updated_by = request.user
        risk.save()
        return redirect("dashboard")

    return render(request, "risks/edit_draft_risk.html", {
        "risk": risk,
        "prob_choices": RiskAssessment.PROBABILITY_CHOICES,
        "impact_choices": RiskAssessment.IMPACT_CHOICES,
    })


# ========= BULK APPROVE =========
@login_required
def bulk_approve_drafts(request):
    if not request.user.is_staff:
        return redirect("dashboard")

    drafts = RiskAssessment.objects.filter(workflow_status='Draft')
    approved_count = drafts.count()
    for risk in drafts:
        risk.description = risk.description.replace("[DRAFT] ", "", 1)
        risk.workflow_status = 'Approved'
        risk.save()

    log_system_event(
        request,
        "approve",
        f"Approved {approved_count} draft risk record(s).",
        metadata={"approved_count": approved_count},
    )

    return redirect("dashboard")
# ========= EXPORT_AND_CLEAR_START =========
@login_required
def export_risks_csv_and_clear(request):
    """
    Staff-only compatibility endpoint.
    Older dashboard buttons used this URL to export and then delete all risks.
    Exporting must not remove saved work, so this now exports only.
    """
    if not request.user.is_staff:
        return redirect("dashboard")

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="risk_register.csv"'
    writer = csv.writer(response)

    writer.writerow([
        'ID', 'Area', 'Description', 'Root Cause', 'Consequences', 'Risk Owner',
        'Inherent Probability', 'Inherent Impact', 'Inherent Rating',
        'Residual Probability', 'Residual Impact', 'Residual Rating'
    ])

    risks_qs = RiskAssessment.objects.all().order_by('-created_at')
    log_system_event(
        request,
        "export",
        f"Exported {risks_qs.count()} risk record(s) through compatibility export endpoint.",
        metadata={"area": "All Departments", "legacy_endpoint": True},
    )
    for risk in risks_qs:
        writer.writerow([
            risk.reference_id,
            risk.area_name,
            risk.description,
            display_level(risk.inherent_probability),
            display_level(risk.inherent_impact),
            display_level(risk.inherent_rating),
            risk.caused_by,
            risk.control_description,
            risk.mitigation_action or "mitigate",
            risk.risk_owner,
            risk.risk_coordinator_name or "Risk & Compliance Officer",
            display_level(risk.residual_rating)
        ])

    return response
# ========= EXPORT_AND_CLEAR_END =========
# ========= CLEAR_RISKS_START =========
@login_required
def clear_all_risks(request):
    """
    Staff-only: delete risks by explicit scope.
    Only works on POST so nobody clears data by visiting a link.
    """
    if not request.user.is_staff:
        return redirect("dashboard")

    if request.method == "POST":
        scope = request.POST.get("scope", "").strip()
        area_name = request.POST.get("area", "").strip()
        confirm_text = request.POST.get("confirm_text", "").strip()
        delete_reason = request.POST.get("delete_reason", "").strip()

        if confirm_text != "DELETE" or not delete_reason:
            return redirect("/?delete_error=confirmation_required")

        if scope == "department" and area_name:
            deleted_refs = list(
                RiskAssessment.objects.filter(area_name=area_name)
                .values_list("reference_id", flat=True)
            )
            deleted_count, _ = RiskAssessment.objects.filter(area_name=area_name).delete()
            log_system_event(
                request,
                "delete",
                f"Deleted {deleted_count} risk record(s) from {area_name}.",
                metadata={
                    "scope": "department",
                    "reason": delete_reason,
                    "deleted_references": deleted_refs,
                },
                area_name=area_name,
            )
            query = urlencode({
                "deleted_area": area_name,
                "deleted_count": deleted_count,
            })
            return redirect(f"/?{query}")

        if scope == "all":
            deleted_refs = list(RiskAssessment.objects.values_list("reference_id", flat=True))
            deleted_count, _ = RiskAssessment.objects.all().delete()
            log_system_event(
                request,
                "delete",
                f"Deleted {deleted_count} risk record(s) from all departments.",
                metadata={
                    "scope": "all",
                    "reason": delete_reason,
                    "deleted_references": deleted_refs,
                },
            )
            query = urlencode({
                "deleted_all": "1",
                "deleted_count": deleted_count,
            })
            return redirect(f"/?{query}")

        return redirect("/?delete_error=missing_scope")

    return redirect("dashboard")
# ========= CLEAR_RISKS_END =========
# ========= BOARD_EXPLANATION_START =========
def _rating_counts(qs, field_name):
    counts = {
        "High": 0,
        "Medium": 0,
        "Low": 0,
    }
    for item in qs:
        value = getattr(item, field_name, "") or ""
        if value in counts:
            counts[value] += 1
    return counts


def _top_risk_themes(risks, limit=5):
    keyword_map = {
        "Fraud / Financial Crime": [
            "fraud", "money laundering", "aml", "cft", "theft",
            "identity theft", "misappropriation", "unauthorized"
        ],
        "Operational Process Breakdown": [
            "process", "delay", "error", "breakdown", "overdue",
            "documentation", "reconciliation", "processing"
        ],
        "Customer / Service Impact": [
            "customer", "complaint", "service", "downtime",
            "reputational", "reputation"
        ],
        "Regulatory / Compliance Exposure": [
            "regulatory", "compliance", "penalty", "sanction",
            "legal", "litigation", "breach"
        ],
        "Technology / Information Security": [
            "system", "it", "ict", "data", "privacy", "breach",
            "access", "security", "cyber", "information leakage"
        ],
        "Credit / Recovery Exposure": [
            "credit", "loan", "recovery", "collections", "default"
        ],
    }

    scores = {k: 0 for k in keyword_map.keys()}

    for risk in risks:
        combined = " ".join([
            risk.description or "",
            risk.caused_by or "",
            risk.consequences or "",
            risk.controls or "",
        ]).lower()

        for theme, words in keyword_map.items():
            if any(word in combined for word in words):
                scores[theme] += 1

    ranked = [(theme, count) for theme, count in scores.items() if count > 0]
    ranked.sort(key=lambda x: (-x[1], x[0]))
    return ranked[:limit]


def _sample_risks(risks, limit=5):
    ranked = sorted(
        risks,
        key=lambda r: (
            {"High": 0, "Medium": 1, "Low": 2}.get(normalize_risk_level(r.residual_rating), 9),
            {"High": 0, "Medium": 1, "Low": 2}.get(normalize_risk_level(r.inherent_rating), 9),
            r.reference_id
        )
    )
    return ranked[:limit]


def _build_board_narrative(area_name, risks):
    total = len(risks)

    if total == 0:
        return {
            "executive_summary": (
                f"No risk records are currently available for {area_name or 'the selected department'}, "
                "so a board-ready explanation cannot yet be generated."
            ),
            "inherent_summary": "No inherent risk profile is available because no risks were found.",
            "residual_summary": "No residual risk profile is available because no risks were found.",
            "control_effectiveness": "Control effectiveness cannot be assessed until risk records are available.",
            "board_recommendation": (
                "Management should ensure the department’s current risk register is populated and validated "
                "before the next board reporting cycle."
            ),
            "top_themes": [],
            "sample_risks": [],
            "inherent_counts": {"High": 0, "Medium": 0, "Low": 0},
            "residual_counts": {"High": 0, "Medium": 0, "Low": 0},
            "improvement_count": 0,
            "unchanged_count": 0,
            "worsened_count": 0,
        }

    inherent_counts = _rating_counts(risks, "inherent_rating")
    residual_counts = _rating_counts(risks, "residual_rating")

    scale = {"Low": 1, "Medium": 2, "High": 3}
    improvement_count = 0
    unchanged_count = 0
    worsened_count = 0

    for risk in risks:
        before = scale.get(risk.inherent_rating, 0)
        after = scale.get(risk.residual_rating, 0)
        if after < before:
            improvement_count += 1
        elif after == before:
            unchanged_count += 1
        else:
            worsened_count += 1

    inherent_high = inherent_counts["High"]
    residual_high = residual_counts["High"]

    area_label = area_name or "Selected Department"

    if inherent_high >= max(1, round(total * 0.5)):
        inherent_tone = (
            "The inherent risk profile is elevated, with a significant share of exposures falling within the "
            "High band before controls are applied."
        )
    elif inherent_high > 0:
        inherent_tone = (
            "The inherent risk profile shows a mixed position, with some material exposures in the higher bands "
            "before controls are applied."
        )
    else:
        inherent_tone = (
            "The inherent risk profile is comparatively contained, with exposures concentrated mainly in the "
            "Medium and Low bands before controls are applied."
        )

    if residual_high == 0:
        residual_tone = (
            "After controls, the residual risk profile appears well contained, with no remaining exposures in the "
            "High band."
        )
    elif residual_high < inherent_high:
        residual_tone = (
            "After controls, the residual risk profile improves relative to the inherent position, although some "
            "higher-risk exposures remain and still require management attention."
        )
    else:
        residual_tone = (
            "After controls, the residual risk profile remains materially elevated, indicating that existing "
            "mitigation measures may not yet be reducing exposure to the desired level."
        )

    if improvement_count >= max(1, round(total * 0.5)):
        effectiveness_text = (
            "Overall, the control environment appears to be having a meaningful moderating effect on risk exposure, "
            "as a majority of risks reduce in rating from inherent to residual position."
        )
    elif improvement_count > 0:
        effectiveness_text = (
            "The control environment is providing partial mitigation benefit, but its impact is uneven across the "
            "department’s risk universe."
        )
    else:
        effectiveness_text = (
            "The current control environment does not yet show clear evidence of risk reduction across the portfolio, "
            "and further strengthening may be required."
        )

    if residual_counts["High"] > 0:
        recommendation = (
            "Board attention is recommended for the remaining High residual exposures. Management should present "
            "targeted remediation actions, named accountabilities, and implementation timelines for those items."
        )
    elif residual_counts["Medium"] > 0:
        recommendation = (
            "The board may note that while controls are reducing exposure, some Medium residual risks remain. "
            "Management should continue focused monitoring and strengthen controls in the affected areas."
        )
    else:
        recommendation = (
            "The board may note that the department’s residual exposure is presently within a more manageable range. "
            "Management should sustain the current control discipline and continue periodic monitoring."
        )

    executive_summary = (
        f"The risk assessment for {area_label} covers {total} identified risk item"
        f"{'' if total == 1 else 's'}. Before controls, {inherent_counts['High']} risk(s) were rated High, "
        f"{inherent_counts['Medium']} Medium, and {inherent_counts['Low']} Low. After accounting for controls, the profile moved to "
        f"{residual_counts['High']} High, {residual_counts['Medium']} Medium, and {residual_counts['Low']} Low. "
        f"This indicates that {improvement_count} risk(s) improved, {unchanged_count} remained unchanged, "
        f"and {worsened_count} worsened between the inherent and residual positions."
    )

    inherent_summary = (
        f"For {area_label}, the inherent risk position reflects the level of exposure that exists before the full "
        f"effect of controls is considered. {inherent_tone} This means the department is naturally exposed to "
        f"operational, compliance, financial, or service-related pressures that could affect performance, customer "
        f"confidence, regulatory standing, or loss outcomes if not actively managed."
    )

    residual_summary = (
        f"The residual risk position reflects the level of exposure that remains after existing controls and response "
        f"measures are considered. {residual_tone} In practical terms, this shows the extent to which current "
        f"controls are helping management contain the department’s most significant risk drivers."
    )

    themes = _top_risk_themes(risks)
    sample_risks = _sample_risks(risks)

    return {
        "executive_summary": executive_summary,
        "inherent_summary": inherent_summary,
        "residual_summary": residual_summary,
        "control_effectiveness": effectiveness_text,
        "board_recommendation": recommendation,
        "top_themes": themes,
        "sample_risks": sample_risks,
        "inherent_counts": inherent_counts,
        "residual_counts": residual_counts,
        "improvement_count": improvement_count,
        "unchanged_count": unchanged_count,
        "worsened_count": worsened_count,
    }


@login_required
def board_explanation(request):
    selected_area = request.GET.get("area", "").strip()
    filter_type = request.GET.get("filter", "approved").strip()

    risks = RiskAssessment.objects.all().order_by("area_name", "reference_id")

    available_areas = list(
        RiskAssessment.objects.exclude(area_name__isnull=True)
        .exclude(area_name__exact="")
        .values_list("area_name", flat=True)
        .distinct()
    )

    if selected_area:
        risks = risks.filter(area_name=selected_area)

    if filter_type == "draft":
        risks = risks.filter(workflow_status='Draft')
    elif filter_type == "approved":
        risks = risks.filter(workflow_status='Approved')
    elif filter_type == "reviewed":
        risks = risks.filter(workflow_status='Reviewed')
    elif filter_type == "closed":
        risks = risks.filter(workflow_status='Closed')

    risk_list = list(risks)
    narrative = _build_board_narrative(selected_area, risk_list)

    context = {
        "selected_area": selected_area,
        "filter_type": filter_type,
        "available_areas": available_areas,
        "risks": risk_list,
        **narrative,
    }
    return render(request, "risks/board_explanation.html", context)
# ========= BOARD_EXPLANATION_END =========
# ========= BOARD_EXPLANATION_START =========
def _rating_counts(qs, field_name):
    counts = {
        "High": 0,
        "Medium": 0,
        "Low": 0,
    }
    for item in qs:
        value = getattr(item, field_name, "") or ""
        if value in counts:
            counts[value] += 1
    return counts


def _top_risk_themes(risks, limit=5):
    keyword_map = {
        "Fraud / Financial Crime": [
            "fraud", "money laundering", "aml", "cft", "theft",
            "identity theft", "misappropriation", "unauthorized"
        ],
        "Operational Process Breakdown": [
            "process", "delay", "error", "breakdown", "overdue",
            "documentation", "reconciliation", "processing"
        ],
        "Customer / Service Impact": [
            "customer", "complaint", "service", "downtime",
            "reputational", "reputation"
        ],
        "Regulatory / Compliance Exposure": [
            "regulatory", "compliance", "penalty", "sanction",
            "legal", "litigation", "breach"
        ],
        "Technology / Information Security": [
            "system", "it", "ict", "data", "privacy", "breach",
            "access", "security", "cyber", "information leakage"
        ],
        "Credit / Recovery Exposure": [
            "credit", "loan", "recovery", "collections", "default"
        ],
    }

    scores = {k: 0 for k in keyword_map.keys()}

    for risk in risks:
        combined = " ".join([
            risk.description or "",
            risk.caused_by or "",
            risk.consequences or "",
            risk.controls or "",
        ]).lower()

        for theme, words in keyword_map.items():
            if any(word in combined for word in words):
                scores[theme] += 1

    ranked = [(theme, count) for theme, count in scores.items() if count > 0]
    ranked.sort(key=lambda x: (-x[1], x[0]))
    return ranked[:limit]


def _sample_risks(risks, limit=5):
    ranked = sorted(
        risks,
        key=lambda r: (
            {"High": 0, "Medium": 1, "Low": 2}.get(normalize_risk_level(r.residual_rating), 9),
            {"High": 0, "Medium": 1, "Low": 2}.get(normalize_risk_level(r.inherent_rating), 9),
            r.reference_id
        )
    )
    return ranked[:limit]


def _build_board_narrative(area_name, risks):
    total = len(risks)

    if total == 0:
        return {
            "executive_summary": (
                f"No risk records are currently available for {area_name or 'the selected department'}, "
                "so a board-ready explanation cannot yet be generated."
            ),
            "inherent_summary": "No inherent risk profile is available because no risks were found.",
            "residual_summary": "No residual risk profile is available because no risks were found.",
            "control_effectiveness": "Control effectiveness cannot be assessed until risk records are available.",
            "board_recommendation": (
                "Management should ensure the department’s current risk register is populated and validated "
                "before the next board reporting cycle."
            ),
            "top_themes": [],
            "sample_risks": [],
            "inherent_counts": {"High": 0, "Medium": 0, "Low": 0},
            "residual_counts": {"High": 0, "Medium": 0, "Low": 0},
            "improvement_count": 0,
            "unchanged_count": 0,
            "worsened_count": 0,
        }

    inherent_counts = _rating_counts(risks, "inherent_rating")
    residual_counts = _rating_counts(risks, "residual_rating")

    scale = {"Low": 1, "Medium": 2, "High": 3}
    improvement_count = 0
    unchanged_count = 0
    worsened_count = 0

    for risk in risks:
        before = scale.get(risk.inherent_rating, 0)
        after = scale.get(risk.residual_rating, 0)
        if after < before:
            improvement_count += 1
        elif after == before:
            unchanged_count += 1
        else:
            worsened_count += 1

    inherent_high = inherent_counts["High"]
    residual_high = residual_counts["High"]

    area_label = area_name or "Selected Department"

    if inherent_high >= max(1, round(total * 0.5)):
        inherent_tone = (
            "The inherent risk profile is elevated, with a significant share of exposures falling within the "
            "High band before controls are applied."
        )
    elif inherent_high > 0:
        inherent_tone = (
            "The inherent risk profile shows a mixed position, with some material exposures in the higher bands "
            "before controls are applied."
        )
    else:
        inherent_tone = (
            "The inherent risk profile is comparatively contained, with exposures concentrated mainly in the "
            "Medium and Low bands before controls are applied."
        )

    if residual_high == 0:
        residual_tone = (
            "After controls, the residual risk profile appears well contained, with no remaining exposures in the "
            "High band."
        )
    elif residual_high < inherent_high:
        residual_tone = (
            "After controls, the residual risk profile improves relative to the inherent position, although some "
            "higher-risk exposures remain and still require management attention."
        )
    else:
        residual_tone = (
            "After controls, the residual risk profile remains materially elevated, indicating that existing "
            "mitigation measures may not yet be reducing exposure to the desired level."
        )

    if improvement_count >= max(1, round(total * 0.5)):
        effectiveness_text = (
            "Overall, the control environment appears to be having a meaningful moderating effect on risk exposure, "
            "as a majority of risks reduce in rating from inherent to residual position."
        )
    elif improvement_count > 0:
        effectiveness_text = (
            "The control environment is providing partial mitigation benefit, but its impact is uneven across the "
            "department’s risk universe."
        )
    else:
        effectiveness_text = (
            "The current control environment does not yet show clear evidence of risk reduction across the portfolio, "
            "and further strengthening may be required."
        )

    if residual_counts["High"] > 0:
        recommendation = (
            "Board attention is recommended for the remaining High residual exposures. Management should present "
            "targeted remediation actions, named accountabilities, and implementation timelines for those items."
        )
    elif residual_counts["Medium"] > 0:
        recommendation = (
            "The board may note that while controls are reducing exposure, some Medium residual risks remain. "
            "Management should continue focused monitoring and strengthen controls in the affected areas."
        )
    else:
        recommendation = (
            "The board may note that the department’s residual exposure is presently within a more manageable range. "
            "Management should sustain the current control discipline and continue periodic monitoring."
        )

    executive_summary = (
        f"The risk assessment for {area_label} covers {total} identified risk item"
        f"{'' if total == 1 else 's'}. Before controls, {inherent_counts['High']} risk(s) were rated High, "
        f"{inherent_counts['Medium']} Medium, and {inherent_counts['Low']} Low. After accounting for controls, the profile moved to "
        f"{residual_counts['High']} High, {residual_counts['Medium']} Medium, and {residual_counts['Low']} Low. "
        f"This indicates that {improvement_count} risk(s) improved, {unchanged_count} remained unchanged, "
        f"and {worsened_count} worsened between the inherent and residual positions."
    )

    inherent_summary = (
        f"For {area_label}, the inherent risk position reflects the level of exposure that exists before the full "
        f"effect of controls is considered. {inherent_tone} This means the department is naturally exposed to "
        f"operational, compliance, financial, or service-related pressures that could affect performance, customer "
        f"confidence, regulatory standing, or loss outcomes if not actively managed."
    )

    residual_summary = (
        f"The residual risk position reflects the level of exposure that remains after existing controls and response "
        f"measures are considered. {residual_tone} In practical terms, this shows the extent to which current "
        f"controls are helping management contain the department’s most significant risk drivers."
    )

    themes = _top_risk_themes(risks)
    sample_risks = _sample_risks(risks)

    return {
        "executive_summary": executive_summary,
        "inherent_summary": inherent_summary,
        "residual_summary": residual_summary,
        "control_effectiveness": effectiveness_text,
        "board_recommendation": recommendation,
        "top_themes": themes,
        "sample_risks": sample_risks,
        "inherent_counts": inherent_counts,
        "residual_counts": residual_counts,
        "improvement_count": improvement_count,
        "unchanged_count": unchanged_count,
        "worsened_count": worsened_count,
    }


@login_required
def board_explanation(request):
    selected_area = request.GET.get("area", "").strip()
    filter_type = request.GET.get("filter", "approved").strip()

    risks = RiskAssessment.objects.all().order_by("area_name", "reference_id")

    available_areas = list(
        RiskAssessment.objects.exclude(area_name__isnull=True)
        .exclude(area_name__exact="")
        .values_list("area_name", flat=True)
        .distinct()
    )

    if selected_area:
        risks = risks.filter(area_name=selected_area)

    if filter_type == "draft":
        risks = risks.filter(workflow_status='Draft')
    elif filter_type == "approved":
        risks = risks.filter(workflow_status='Approved')
    elif filter_type == "reviewed":
        risks = risks.filter(workflow_status='Reviewed')
    elif filter_type == "closed":
        risks = risks.filter(workflow_status='Closed')

    risk_list = list(risks)
    narrative = _build_board_narrative(selected_area, risk_list)

    context = {
        "selected_area": selected_area,
        "filter_type": filter_type,
        "available_areas": available_areas,
        "risks": risk_list,
        **narrative,
    }
    return render(request, "risks/board_explanation.html", context)
# ========= BOARD_EXPLANATION_END =========
